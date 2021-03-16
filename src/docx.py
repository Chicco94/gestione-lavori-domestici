from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from docx.text.run import Font

from docx.shared import Inches,RGBColor

from datetime import datetime
from src.help_variables import giorni_settimana



def prepare_document_styles(document):
	# WD_STYLE_TYPE.CHARACTER = 2 Se metto l'elemento dell'enum 
	# lancia un errore in editor che mi da fastidio
	style = document.styles.add_style('mamma',2)
	style.font.bold = True
	style.font.color.rgb = RGBColor(0xff, 0x00, 0xff)

	style = document.styles.add_style('papà',2)
	style.font.bold = True
	style.font.color.rgb = RGBColor(0x00, 0x00, 0xff)

	style = document.styles.add_style('alberto',2)
	style.font.bold = True
	style.font.color.rgb = RGBColor(0x00, 0xff, 0x00)

	style = document.styles.add_style('enrico',2)
	style.font.bold = True
	style.font.color.rgb = RGBColor(0xdc, 0x14, 0x3c)

	style = document.styles.add_style('carlo',2)
	style.font.bold = True
	style.font.color.rgb = RGBColor(0xff, 0xd7, 0x00)

	style = document.styles.add_style('samantha',2)
	style.font.bold = True
	style.font.color.rgb = RGBColor(0xff, 0x69, 0xb4)

	style = document.styles.add_style('giulia',2)
	style.font.bold = True
	style.font.color.rgb = RGBColor(0xff, 0x14, 0x93)

	return document


def create_cell_content(cell,header, day, styles):
	cell.paragraphs[0].text = header
	cell.paragraphs[0].style = 'Header'
	for job,users in day.items():
		par = cell.add_paragraph("{} svolto da ".format(job), style='ListBullet')
		for user in users:
			par.add_run(user,style=styles[user])


def create_schedule_document(schedule,footer):
	'''Save the schedule in a word file'''
	document = prepare_document_styles(Document())
	document.add_heading('Lavori per la settimana del {}'.format(datetime.now().strftime('%d-%m-%Y')), 0)

	for index_of_day in range(0,7):
		document.add_heading(giorni_settimana[index_of_day], level=1)

		table = document.add_table(rows=0, cols=2)
		row_cells = table.add_row().cells
		create_cell_content(row_cells[0],'Mattina',schedule[index_of_day],document.styles)
		create_cell_content(row_cells[1],'Pranzo',schedule[index_of_day+7],document.styles)
		row_cells = table.add_row().cells
		create_cell_content(row_cells[0],'Pomeriggio',schedule[index_of_day+14],document.styles)
		create_cell_content(row_cells[1],'Sera',schedule[index_of_day+21],document.styles)
	
	document.add_paragraph(footer)

	filename = datetime.now().strftime('%Y-%m-%d')
	document.save('./build/{}.docx'.format(filename))
	return True